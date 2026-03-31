# -*- coding: utf-8 -*-
"""One-off: generate investor analysis Word doc. Run: python docs/build_investor_docx.py from crm-mvp."""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).resolve().parent / "Polden_Investor_Business_and_Technical_Audit.docx"


def add_h(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_p(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p


def add_bullets(doc, items):
    for t in items:
        doc.add_paragraph(t, style="List Bullet")


def main():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Inches(0.9)
    s.bottom_margin = Inches(0.9)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Полдень (Polden)\n")
    r.bold = True
    r.font.size = Pt(18)
    r2 = t.add_run("Полный анализ продукта и технический аудит кодовой базы\n")
    r2.font.size = Pt(14)
    r3 = t.add_run("Материал для инвесторов\n")
    r3.font.size = Pt(12)
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m = meta.add_run(
        "Источник: авторитетный контур репозитория CRM/crm-mvp, landing-order.\n"
        "Метод: обзор README, Prisma schema, маршрутов API (server.js), документации docs/.\n"
        "Дата подготовки: март 2026.\n\n"
        "Оговорка: финансовые показатели и юридический due diligence в документ не включены — "
        "только описание продукта по коду и оценка зрелости платформы."
    )
    m.font.size = Pt(9)
    m.italic = True

    add_h(doc, "1. Резюме для инвестора", 1)
    add_p(
        doc,
        "«Полдень» — это интегрированный контур для доставки обедов (или аналогичного формата): "
        "публичный заказ через лендинг, внутренняя CRM на React, единый backend на Node.js/Express с Prisma. "
        "В кодовой базе зафиксированы: живой путь заказа (публичный API → БД → просмотр в CRM), "
        "редактирование меню на день по филиалам, атрибуция заказов (UTM/путь), KPI запуска, "
        "контент-пайплайн под VK, лид-форма VK (лиды отдельно от заказов), а также расширенный модуль "
        "«Kitchen Economics» (себестоимость, склад, закупки, списания) со спецификациями в docs/. "
        "Технически проект позиционируется как MVP с явной траекторией развития back-office и маркетинга.",
    )

    add_h(doc, "2. Описание бизнеса (как следует из продукта)", 1)
    add_h(doc, "2.1. Ценностное предложение", 2)
    add_p(
        doc,
        "Автоматизация цикла «меню на завтра → приём заказа на сайте → учёт в CRM по филиалу и дате доставки» "
        "с возможностью отслеживать эффективность запусков (KPI, источники) и готовить контентные кампании под VK. "
        "Оператор видит заказы в одном интерфейсе; кухня и закупки — в отдельных API/UI слоях (часть в спецификациях v1).",
    )
    add_h(doc, "2.2. Пользователи и сценарии", 2)
    add_bullets(
        doc,
        [
            "Конечный клиент — оформляет заказ на публичном лендинге (landing-order), выбирая филиал и позиции из меню на день.",
            "Оператор/менеджер — в CRM задаёт меню на дату, просматривает заказы по филиалу и дате доставки, работает с KPI.",
            "Маркетинг — планирование контента VK, UTM, drill-записи (аудит связи контент → заказ).",
            "VK-бот — собирает лиды (модель VkLead), не подменяя собой полноценный заказ; оператор оформляет DeliveryOrder отдельно.",
            "Back-office (кухня/склад) — сущности Ingredient, StockMovement, PurchaseDraft, ProductionWriteoff и др. по schema.prisma и docs/KITCHEN_ECONOMICS_*.md.",
        ],
    )
    add_h(doc, "2.3. Ключевые бизнес-процессы в системе", 2)
    add_bullets(
        doc,
        [
            "Публикация меню: MenuDayItem привязан к Branch + date + position (слоты 1–10 на лендинге).",
            "Публичная витрина: GET /api/public/menu-day, GET /api/public/branches; расчёт quote и создание заказа: POST /api/public/delivery-orders*.",
            "Исполнение: DeliveryOrder с deliveryDate, позициями DeliveryOrderItem, опционально attributionJson.",
            "Запуски: dashboard launch-kpis, ContentItem + LaunchDrillRecord для ручного аудита цепочки.",
        ],
    )

    add_h(doc, "3. Архитектура и стек (аудит)", 1)
    add_h(doc, "3.1. Компоненты", 2)
    add_bullets(
        doc,
        [
            "Backend: CRM/crm-mvp/backend — Express, Prisma ORM, по умолчанию SQLite (DATABASE_URL), порт разработки 4000.",
            "Frontend CRM: CRM/crm-mvp/frontend — Vite + React, токен X-CRM-Token для защищённых маршрутов.",
            "Публичный заказ: landing-order — статический/лёгкий фронт, обращается к публичному API.",
            "Верификация релиза: npm run verify:launch из корня crm-mvp (скрипт scripts/verify-launch.mjs).",
        ],
    )
    add_h(doc, "3.2. API (обзор по server.js)", 2)
    add_p(doc, "Публичные (без CRM-токена): /health, /api/public/branches, /api/public/menu-day, POST quote и delivery-orders.")
    add_p(doc, "Защищённые (requireCrmToken): delivery-orders список, menu-day-items, upsert меню, launch-kpis, VK-bot readiness, content-pipeline, procurement-board, economics/production/purchase APIs и др.")
    add_h(doc, "3.3. Данные", 2)
    add_p(
        doc,
        "Prisma-схема объединяет заказы (Branch, MenuDayItem, DeliveryOrder), маркетинг (ContentItem, LaunchDrillRecord), "
        "VK (VkConversationState, VkLead) и кухонно-складской контур (Unit, Ingredient, Dish, StockMovement, PurchaseDraft, …). "
        "SQLite удобен для MVP и пилотов; для масштаба инвестору стоит планировать миграцию на управляемую СУБД и резервное копирование.",
    )

    add_h(doc, "4. Зрелость и операционная готовность", 1)
    add_bullets(
        doc,
        [
            "Задокументирован ежедневный контур запуска: docs/LAUNCH_BASELINE_HANDOFF.md, OPERATOR_RUNBOOK_LAUNCH.md — снижает риск человеческой ошибки (филиал, дата, API).",
            "Старт backend: проверка пути БД (databaseEnv.js), опционально строгий режим без пустых Branch, расширенный /health.",
            "Риски эксплуатации: один токен CRM_INTERNAL_TOKEN для внутренних API — нужна ротация и разделение сред (dev/stage/prod).",
            "Публичные эндпоинты должны быть защищены rate limiting и WAF на уровне инфраструктуры (в коде не детализировано).",
        ],
    )

    add_h(doc, "5. Модули «вне ядра заказа» (оценка для инвестора)", 1)
    add_bullets(
        doc,
        [
            "Kitchen Economics v1 — описан в англоязычных спецификациях (KITCHEN_ECONOMICS_V1_SPEC.md и связанные); в схеме БД присутствуют сущности — глубина внедрения в UI следует проверять отдельно.",
            "Контент-пайплайн и VK drill — модели и API есть; это инструмент дисциплины маркетинга, не автоматическая оптимизация.",
            "VK лиды — отдельная сущность от заказа; конверсия лида в заказ остаётся на операторе — важно для оценки unit-экономики процесса.",
        ],
    )

    add_h(doc, "6. Риски и пробелы (честный список)", 1)
    add_bullets(
        doc,
        [
            "Масштабирование: SQLite как дефолт ограничивает одновременную запись и HA-сценарии.",
            "Безопасность: единый shared secret для CRM API; необходимы политики секретов, HTTPS, ограничение CORS в проде.",
            "Дублирование каталогов в workspace (пути new/, дубликаты crm-mvp) — для инвестиций важно закрепить один authoritative репозиторий.",
            "Финансовая и юридическая сторона, персональные данные (152-ФЗ/GDPR) — требуют отдельного due diligence, не отражены в коде.",
        ],
    )

    add_h(doc, "7. Вывод", 1)
    add_p(
        doc,
        "Продукт представляет собой связный MVP доставки с публичным каналом продаж, операторской CRM и заделом под маркетинг (VK) "
        "и углублённый операционный учёт кухни. Ядро заказа технически прослеживается от лендинга до модели DeliveryOrder. "
        "Для инвестиционного решения рекомендуется дополнить этот технический аудит финансовой моделью, метриками удержания клиентов B2B (если продажа ресторанам) и планом миграции данных/инфраструктуры.",
    )

    doc.save(OUT)
    print("Wrote", OUT)


if __name__ == "__main__":
    main()
